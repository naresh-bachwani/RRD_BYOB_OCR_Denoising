from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
import pdfminer
import os
import re
import autocorrect
from IPython.display import clear_output
global base_path
import sys
base_path = 'C:/Users/Naresh/Desktop/RRD/DataSet_BYOBChallenge2_2'

def extract_layout_by_page(pdf_path):
    """
    Extracts LTPage objects from a pdf file.
    
    slightly modified from
    https://euske.github.io/pdfminer/programming.html
    """
    laparams = LAParams()

    fp = open(pdf_path, 'rb')
    parser = PDFParser(fp)
    document = PDFDocument(parser)

    if not document.is_extractable:
        raise PDFTextExtractionNotAllowed

    rsrcmgr = PDFResourceManager()
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)

    layouts = []
    for page in PDFPage.create_pages(document):
        interpreter.process_page(page)
        layouts.append(device.get_result())

    return layouts


# page_layouts = extract_layout_by_page(my_file)

def flatten(lst):
    """Flattens a list of lists"""
    return [subelem for elem in lst for subelem in elem]

def extract_characters(element):
    """
    Recursively extracts individual characters from 
    textw elements. 
    """
    if isinstance(element, pdfminer.layout.LTChar):
        return [element]

    if any(isinstance(element, i) for i in TEXT_ELEMENTS):
        return flatten([extract_characters(e) for e in element])

    if isinstance(element, list):
        return flatten([extract_characters(l) for l in element])

    return []

'''rounds offf the y coordinate'''
def roundy(characters):
    chart=characters[0].bbox[1]
    for char in characters:
    
        chare=list(char.bbox)
        tmp = int(chare[1])
        if(abs(tmp-chart)<5):
            tmp=chart
        chare[1] = tmp
        chare=tuple(chare)
        char.bbox=chare
        chart=tmp

'''create chunks'''
def createchunks(characters,bigchunk,linechunk,chunk):
    lchar = characters[0].bbox[1] #last char y1
    lchar2=characters[0].bbox[2] #last char x2
    ltext = characters[0].get_text()
    # horizontal_dist = max(hdist1)-min(hdist0)

    # q=horizontal_dist/88

    ''' generates all the respective chunks'''
    flag=0
    for char in characters:

        if(char.bbox[1]==lchar):
            flag=1
            if((char.bbox[0]-lchar2)<8):
                chunk.append(char)
            else:
                linechunk.append(chunk)
                chunk=[]
                chunk.append(char)
        else:
            flag=0
            linechunk.append(chunk)

            bigchunk.append(linechunk)
            linechunk=[]
            chunk=[]
            chunk.append(char)
        lchar=char.bbox[1]
        if(char.get_text()!=' '):
            lchar2=char.bbox[2]
        ltext=char.get_text()
    if(flag==1):
            flag=0
            linechunk.append(chunk)

            bigchunk.append(linechunk)
            linechunk=[]
            chunk=[]
            chunk.append(char)

'''find the starting and  ending coordinate of a chunk'''
def coordinates(bigchunk,chunkstart,chunkend):
    for linechunk in bigchunk:
        for chunk in linechunk:
            start=[]
            end=[]
            for c in chunk:
                if(c.get_text() !='_'):
                    start.append(c.bbox[0])
                if(c.get_text()!=' ' and c.get_text()!='_'):
                    end.append(c.bbox[2])
   
            if(len(start)>0):
                chunkstart.append(min(start))
                
            else:
                chunkstart.append(0)
                
            if(len(end)>0):
                chunkend.append(max(end))
                
            else:
                chunkend.append(0)
            

'''rounding of starting coordinate so that slight error in retrieving coordinates is dealtwith'''
def roundcoordinates(chunkstart):
    for i in range(len(chunkstart)):
        for j in range(i+1,len(chunkstart)):
            if(abs(chunkstart[j]-chunkstart[i])<12):
                chunkstart[j]=min(chunkstart[i],chunkstart[j])
                if(chunkstart[i]!=chunkstart[j]):
                    chunkstart = [chunkstart[j] if x==chunkstart[i] else x for x in chunkstart] 
    return chunkstart
    #         if(abs(chunkend[j]-chunkend[i])<q*4):
    #             chunkend[j]=chunkend[i]

'''gnerate a indexchunk which stores multiple information about the chunk'''
''' it is in form of a dictionary
    { start:
      end:
      text:
      lineno:
      y:
      isTable:
    } '''
def genindex(bigchunk,indexchunk):
    count=0
    lineno=0
    for linechunk in bigchunk:
        lineno+=1
        for chunk in linechunk:
            isTable=False
            tex = [c.get_text() for c in chunk]
            tex = ''.join(tex)
            tex = tex.lstrip()
            tex = tex.rstrip()
            new = {
                'start':chunkstart[count],
                'end': chunkend[count],
                'text':tex,
                'lineno':lineno,
                'y': (chunk[0].y0+chunk[0].y1)/2,
                'isTable': isTable
            }
    #         new=[chunkstart[count],isTable,tex,lineno,False,(chunk[0].y0+chunk[0].y1)/2,chunkend[count]]
            indexchunk.append(new)
            count+=1

'''removes any chunk which do not contain alphas or numbers or currency symbol'''
def remove_error(indexchunk):
    toremove=[]
    for char in indexchunk:
        if (('£'not in char['text']) and ('$' not in char['text'])):
            if char['text']=='z-':
                toremove.append([char])
            if(len(char['text'])==1 and not(char['text'].isdigit())):
                toremove.append([char])
            elif ((not re.search('\w',char['text']))):
                toremove.append([char])
            elif(char['start']==char['end']):
                toremove.append([char])
    for i in toremove:
        indexchunk.remove(i[0])
        

def table_feat(indexchunk):
    '''add chunks to line_indchunk'''
    ####################################################################################
    ini = indexchunk[0]['lineno']
    line_indchunk=[]
    remotechunk=[]
    for chunk in indexchunk:
        if(chunk['lineno']==ini):
            remotechunk.append(chunk)
        else:
            line_indchunk.append(remotechunk)

            ini=chunk['lineno']
            remotechunk=[]
            remotechunk.append(chunk)
    line_indchunk.append(remotechunk)
    ####################################################################################

    '''check whether a chunk is part of table or not'''
    #***********************************************************************************
    count=0
    for remotechunk in line_indchunk:
        if(len(remotechunk)>1):
            isTable=False
            for chun in remotechunk:
                isTable=False
                flag=0
                for chunk in remotechunk:

                    tex=chunk['text']
                    if(tex!=' '):
                        isTable=True
                    if (not(re.search('[a-zA-Z]', tex))):
                        flag=1
                if(flag==0):
                    isTable=False
                indexchunk[count]['isTable']=isTable
                count+=1
        else:
            count+=1
    #***********************************************************************************
#     indexchunk = serial_corr(indexchunk)
            
    '''some minor changes to table elements'''
    ####################################################################################
    lin=[]
    l=[]
    line=1
    for chunk in indexchunk:

        if chunk['isTable'] ==True:

            if line!=chunk['lineno']:
                line=chunk['lineno']
                if(l!=[]):
                    lin.append(l)
                l=[]

        else:
            line=chunk['lineno']
            if(l!=[]):
                lin.append(l)
            l=[]
        l.append(chunk)
    if l!=[]:
        lin.append(l)
    tochange=[]
    for i in range(len(lin)):
        if len(lin[i])==2:
            if i<len(lin)-1:
                if not (lin[i-1][0]['isTable'] and lin[i+1][0]['isTable']):
                    lin[i][0]['text'].replace(".","");
                    if not re.match('/d', lin[i][0]['text'], flags=re.IGNORECASE):
                        if re.match('[a-zA-Z]',lin[i][1]['text']):
                            tochange.append(lin[i][0]['lineno'])
                        
            elif i==0:
                 if not ( lin[i+1][0]['isTable']):
                    lin[i][0]['text'].replace(".","");
                    if not re.match('/d', lin[i][0]['text'], flags=re.IGNORECASE):
                        if re.match('[a-zA-Z]',lin[i][1]['text']):
                            tochange.append(lin[i][0]['lineno'])
                       
            else:
                if not (lin[i-1][0]['isTable']):
                    lin[i][0]['text'].replace(".","");
                    if not re.match('/d', lin[i][0]['text'], flags=re.IGNORECASE):
                        if re.match('[a-zA-Z]',lin[i][1]['text']):
                            tochange.append(lin[i][0]['lineno'])
        
    
    for chunk in indexchunk:
        if chunk['lineno'] in tochange:
            chunk['isTable']=False
   ######################################################################################
    
    
    
    
    
    '''add more elements to table i.e element whichwere left out'''
    for i in range(1,len(indexchunk)-1):
    
        if((indexchunk[i+1]['isTable'])==True and (indexchunk[i-1]['isTable'])==True and indexchunk[i]['isTable']==False):
            indexchunk[i]['isTable']=True

'''create a range list which contains range of x coordinate for each column'''
def gen_range(tablechunk):
    ranges=[]
    if(len(tablechunk)>0):
        ranges.append([tablechunk[0]['start'],tablechunk[0]['end']])
    for chunk in tablechunk:
        flag=0
        for ran in ranges:
    
            if ((chunk['start']<=ran[1]) and (chunk['end']>=ran[0])):

                ran[0]=min(chunk['start'],ran[0])
                ran[1]=max(ran[1],chunk['end'])
    
                flag=1

    
        if(flag==0):

            ranges.append([chunk['start'],chunk['end']])
            ranges=sorted(ranges,key=lambda x:x[0])
    
    return ranges

'''assign column to each chunk'''
def assign_col(tablechunk):
    for chunk in tablechunk:
    
        for i in range(len(ranges)):
            if(chunk['start']>=ranges[i][0] and chunk['end']<=ranges[i][1]):
                chunk['col']=i

'''assign row number to each chunk'''
def assign_row(tablechunk):
    ini = tablechunk[0]['lineno']
    fin = tablechunk[-1]['lineno']
    nrows = fin-ini+1
    row=0
    tablechunk[0]['row'] = row
    for i in range(1,len(tablechunk)):
        if(tablechunk[i]['lineno']!=tablechunk[i-1]['lineno']):
            row+=1
        tablechunk[i]['row'] = row

'''craeate a final chunk having both normal and table chunk'''
def create_finalchunk(indexchunk,tablechunk,finalchunk):
    i=0
    ntables=0
    count=0
    for chunk in indexchunk:
        if(chunk['isTable']==True):
            if(count==0):
                ntables+=1
                count=1
            finalchunk.append(tablechunk[i])
            i+=1
        else:
            count=0
            finalchunk.append(chunk)
    return ntables

'''Returns the no of  rows of table'''
def dimension(finalchunk):
    dimensions=[]
    row=0
    flag=0
    line=0
    for chunk in finalchunk:
        if chunk['isTable'] :
            
            if line!=chunk['lineno']:
                row+=1
                line=chunk['lineno']
            chunk['row']=row-1

        else:
            if row ==0:
                line=chunk['lineno']

            else:
                dimensions.append(row)
                row=0
    if row ==0:
        line=chunk['lineno']

    else:
        dimensions.append(row)
        row=0
    return dimensions


'''creates the table'''

def create_doc(document_chunk):
    for finalchunk in document_chunk:
        count=0
        tab = 0
        lchunk = finalchunk[0][0]
        tex=''
        for chunk in finalchunk[0]:
            if(chunk['isTable']==False):
                '''Arrange  te free flowing text'''
                
                if(abs(lchunk['y']-chunk['y'])>15 or (abs(lchunk['start']-chunk['start'])>10 and lchunk['lineno']!=chunk['lineno'])):
                    
                    
                    document.add_paragraph(tex)
                    tex=''
                space=1
                if(tex==''):
                    space=int((chunk['start']-80)/(5.5)*1.4)

                for i in range(space):
                    tex = tex+' '
                tex+=chunk['text']
                
                
                count=0
                if(chunk['lineno']!=lchunk['lineno']):
                    lchunk=chunk
            else:
                '''Arrange the tables'''
                if(tex!=''):
                    document.add_paragraph(tex)
                    tex=''
                if(count==0):
                    a = finalchunk[1][tab]
                    table = document.add_table(cols=finalchunk[2],rows=finalchunk[1][tab])
                    tab+=1
                    table.autofit=False
                    count=1
        
                i = chunk['row']
                j = chunk['col']
                
                cell = table.cell(i,j)
                cell.text=chunk['text']
        if(tex!=''):
            document.add_paragraph(tex)
            tex=''
        document.add_page_break()


replace={
    '0':['o'],
    '1':['i'],
    '5':['S'],
    '6':['b'],
    '8':['B'],
    'b':['6'],
    'B':['8'],
    'C':['6'],
    'D':['0'],
    'G':['6'],
    'I':['1'],
    'i':['1'],
    'l':['1'],
    'o':['0'],
    'O':['0'],
    'Q':['0'],
    's':['5'],
    'S':['5']
}

'''correction of text'''

def correction(finalchunk):
    for chunk in finalchunk:
        text=chunk['text']
        line =text.split()
        text=""
        for word in line:
            if not len(re.findall('[0-9-.,$£_]',word))==len(word):
                if not len(re.findall('[A-Z]',word))==len(word):
                    if len(re.findall('[A-z]',word))==len(word):
                        word=autocorrect.spell(word)
    #                     print("onlyletter ",word)
                    else:
                        digit=sum(c.isdigit() for c in word)
                        alpha= sum(c.isalpha() for c in word)
                        others  = len(word)-digit-alpha

    #                     print(digit,alpha,others)
                        if others>=(alpha+digit) or others>=3:
                            word=""
                        elif digit>alpha:
                            for c in word:
                                if c.isalpha():
                                    if c in replace.keys():
                                        c=replace[c][0]
    #                         print("digitmore  ",word)
                        else:
                            for c in word:
                                if c.isdigit():
                                    if c in replace.keys():
                                        c=replace[c][0]
                            word=autocorrect.spell(word)
            else:
                word=word.replace("_","")
            if word!="":
                text=" ".join([text,word])
    #              print("nomore  ",word)
    #             else:
    #                 print ("name " ,word)
    #         else:

        chunk['text']=text
    return finalchunk
'''main loop'''


directoryfiles=os.listdir(base_path) #file in the given directory
files=[]
i=0
for filename in directoryfiles:#loop in files
    if('.pdf' not in filename):
        continue
    files.append(filename)
    filee = filename.rstrip('.pdf')

    my_file = os.path.join(base_path + "/" + filee+'.pdf')
   
    document_chunk=[]
    
    page_layouts = extract_layout_by_page(my_file)
    len(page_layouts)
    objects_on_page = set(type(o) for o in page_layouts[1])
    objects_on_page

    for pageno in range(len(page_layouts)): #loop in every page
        text1=[]
        TEXT_ELEMENTS = [
            pdfminer.layout.LTTextBox,
            pdfminer.layout.LTTextBoxHorizontal,
            pdfminer.layout.LTTextLine,
            pdfminer.layout.LTTextLineHorizontal
        ]

        current_page = page_layouts[pageno]

        texts = []
        rects = []

        # seperate text and rectangle elements
        for e in current_page:
            if isinstance(e, pdfminer.layout.LTTextBoxHorizontal):
                texts.append(e)
            elif isinstance(e, pdfminer.layout.LTFigure):
                rects.append(e)

        # sort them into 
        characters = extract_characters(texts)

        #rounds of the y coordinate
        roundy(characters)

        
        #sort the characters in decreasing order of y followed by increasing order of x coordinate
        characters=sorted(characters, key=lambda element: (-element.bbox[1], element.bbox[0]))

        #rounds of the y coordinate
        chart=characters[0].bbox[1]
        for char in characters:
            chare=list(char.bbox)
            tmp = (chare[1])
            if(abs(tmp-chart)<5):

                tmp=chart
            chare[1] = tmp
            chare=tuple(chare)
            char.bbox=chare
            chart=tmp
        #again sort the elements
        characters=sorted(characters, key=lambda element: (-element.bbox[1], element.bbox[0]))


        '''Now make chunks'''
        bigchunk = [] #stores all the line chunks
        linechunk = [] #stores all the chunks i.e chunk of words cluster in chunks
        chunk = [] #smallest block of characters
        createchunks(characters,bigchunk,linechunk,chunk)

        '''find the starting and ending points of every chunk '''
        chunkstart=[]
        chunkend=[]
        coordinates(bigchunk,chunkstart,chunkend)

        '''round the coordinate to deal with slight errors'''
        chunkstart = roundcoordinates(chunkstart)
        
        '''gnerate a indexchunk which stores multiple information about the chunk'''
        ''' it is in form of a dictionary
            { start:
              end:
              text:
              lineno:
              y:
              isTable:False(default)
            } '''
        indexchunk=[]
        genindex(bigchunk,indexchunk)

        '''removes any chunk which do not contain alphas or numbers or currency symbol'''
        remove_error(indexchunk)

        '''Assign isTable feature to each chunk'''
        table_feat(indexchunk)


        '''create a new table chunk which contains elements for table'''
        tablechunk=[]
        for chunkq in indexchunk:
            if(chunkq['isTable']==True):
                tablechunk.append(chunkq)


        

        '''store the range of each column of table in form of coordinates'''
        ranges = gen_range(tablechunk)
        ncols = len(ranges) 
        
        '''assign col no to each chunk'''
        assign_col(tablechunk)

        

        if(len(tablechunk)>0):
            '''assign row no to each chunk'''
            assign_row(tablechunk)

        
        '''create final chunk'''
        finalchunk=[]
        ntables = create_finalchunk(indexchunk,tablechunk,finalchunk)

        ntables

        finalchunk=correction(finalchunk)
        '''dimensions stores the row number of each table'''
        dimensions = dimension(finalchunk)

        '''generate document chunk'''
        '''(finalchunk,dimensions,ncols)'''
        document_chunk.append((finalchunk,dimensions,ncols))

    from docx import Document
    document = Document()
    
    ''' create the doument'''
    create_doc(document_chunk)
     
    document.save(filee+'.docx')#save the file
    
    print(filee)
    print(i)
    i=i+1
    os.system('cls')
    

directoryfiles

len(files)

